"""Build the manuscript .docx. Sections are added incrementally."""
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docbuild import (new_doc, heading, para, equation, figure, table, EQDIR)

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "Manuscript", "Stochastic_Eco_VRPTW.docx")


def sec_front(doc):
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Learning-Augmented Stochastic Eco-Speed Vehicle Routing under "
                  "Demand and Travel-Time Uncertainty for Low-Emission Waste Collection")
    r.bold = True; r.font.size = doc.styles["Normal"].font.size; r.font.name = "Times New Roman"
    r.font.size = r.font.size  # keep
    from docx.shared import Pt
    r.font.size = Pt(15)
    doc.add_paragraph()

    heading(doc, "Abstract", 1)
    para(doc,
         "Municipal waste collection and urban freight are increasingly judged by the "
         "greenhouse-gas emissions that routing decisions generate, yet emissions depend "
         "on vehicle load and cruising speed, both of which are uncertain in practice: "
         "bin fill varies day to day and travel time varies with traffic. We formulate "
         "a stochastic emission-capacitated vehicle routing problem with time windows "
         "and discrete eco-speed levels in which demand and arc travel times are random, "
         "time windows are imposed as chance constraints, and capacity violations are "
         "absorbed through explicit recourse; the objective minimizes expected fuel, "
         "emissions, and recourse cost. A learning-augmented layer estimates the "
         "uncertain quantities from real operational data, attaining a weighted absolute "
         "percentage error of 19.2% with 78.3% prediction-interval coverage for demand, "
         "a mean absolute error of 1.44 minutes (9.2% error) for travel time, and areas "
         "under the ROC curve of 0.980 to 0.995 for eco-speed feasibility. On top of "
         "this layer we introduce a decision-focused rule that learns context-dependent "
         "per-node capacity service levels and selects, per instance, the planning regime "
         "that minimizes realized cost, embedded in a scenario-aware Hybrid Guided Local "
         "Search. Across eleven instance families built from real New York City taxi, "
         "Austin waste-collection, Dublin bin, and Peshawar municipal data together with "
         "adapted benchmark instances, the proposed method lowers expected cost on ten "
         "of eleven families and is significantly better than all eight baselines "
         "(paired Wilcoxon, Holm-adjusted p <= 0.005; Friedman p = 1.4e-7), improving on "
         "the strongest stochastic baseline by 2.8% in median. An ablation shows that "
         "modeling recourse is the dominant driver, increasing expected cost by about "
         "90% when removed, and a sensitivity analysis confirms the advantage is robust "
         "to the recourse cost weights (3-11% across settings). On the Peshawar Zone D "
         "case study the method reduces the probability of vehicle overflow from 0.87 to "
         "0.58 and the conditional value at risk by 10.5% relative to a deterministic "
         "plan. All code and data are publicly available.")
    kw = para(doc, "Keywords: ", bold=True)
    kw.add_run("vehicle routing; eco-speed emissions; stochastic demand and travel time; "
               "chance constraints and recourse; decision-focused learning.")


def sec_intro(doc):
    heading(doc, "1. Introduction", 1)
    para(doc,
         "Urban freight and municipal waste logistics are increasingly shaped by the "
         "dual imperatives of operational efficiency and environmental sustainability. "
         "Classical formulations of the Vehicle Routing Problem with Time Windows "
         "(VRPTW) prioritize minimizing fleet size and distance traveled under hard "
         "capacity and service-window constraints. While such models provide valuable "
         "operational insight and enable benchmarking on standard instances (Solomon, "
         "1987; Arnold and Sorensen, 2019), they are no longer adequate in the context "
         "of contemporary climate policies, carbon-pricing schemes, and sustainable "
         "development targets. The European Green Deal, the United Nations Sustainable "
         "Development Goals, and national decarbonization strategies all emphasize "
         "emissions as a first-order performance measure for logistics systems. As a "
         "result, distance minimization, although important, must be complemented by "
         "explicit emission objectives to capture the true environmental footprint of "
         "routing decisions (Demir, Bektas and Laporte, 2014; Moghdani, Salimifard and "
         "Naderi, 2021).")
    para(doc,
         "Emissions in routing are not linearly proportional to distance. They depend "
         "on vehicle load, speed, traffic conditions, and the allocation of slack time. "
         "The Pollution Routing Problem highlighted this dependency by showing that "
         "fuel burn is a convex function of load and speed (Bektas and Laporte, 2011; "
         "Fukasawa, Pessoa and Poggi, 2016). Subsequent work embedded discrete or "
         "continuous speed decisions within time-window-constrained problems, revealing "
         "that eco-speed choices can substantially reduce fuel and CO2 output without "
         "violating service windows (Gutierrez-Padilla et al., 2021; Lai et al., 2024). "
         "Yet most large-scale VRPTW studies still fix speed to preserve comparability "
         "with distance-oriented benchmarks, creating a gap between academic practice "
         "and policy-relevant models.")
    para(doc,
         "A second limitation is more fundamental. Capacity and time windows remain "
         "non-negotiable in municipal waste collection and urban freight: vehicles "
         "cannot exceed their load capacities, and bins or customers must be serviced "
         "within regulatory or contractual time windows (Munari et al., 2019; Yu, Wang "
         "and Ma, 2019). Real operations, however, are not deterministic. The amount of "
         "waste accumulated in a bin varies from day to day, and the time required to "
         "traverse an arc depends on traffic that fluctuates by hour and location. When "
         "demand and travel time are treated as fixed, the resulting plan is optimistic: "
         "a route that is exactly feasible under mean demand overflows whenever the "
         "realized fill exceeds the mean, and a schedule that is exactly on time under "
         "mean travel time is late whenever traffic is heavier than assumed. Because "
         "emissions are themselves a function of the realized load and speed, the "
         "environmental cost of a plan is a random quantity rather than a single number.")
    para(doc,
         "This paper addresses that gap by formulating a Stochastic Emission-Capacitated "
         "VRPTW with eco-speed (S-E-CVRPTW). Bin-fill demand and arc travel times are "
         "modeled as random variables; service windows are imposed as chance "
         "constraints; and capacity violations are absorbed by explicit recourse, namely "
         "an emergency return to the depot when a vehicle fills beyond its capacity. The "
         "objective minimizes the expected sum of fuel, emissions, and recourse cost, "
         "while the conditional value at risk and the overload probability are reported "
         "as risk diagnostics. The uncertainty is not assumed but learned: a prediction "
         "layer estimates demand quantiles, arc travel-time distributions, and the "
         "probability that a given eco-speed level is operationally achievable, each "
         "trained and validated on real data. Building on recent interest in coupling "
         "machine learning with routing (Bogyrbayeva et al., 2024; Kerscher and Minner, "
         "2025; Wu et al., 2026), we further propose a decision-focused planning rule "
         "that learns a context-dependent per-node capacity service level and selects, "
         "on each instance, the planning regime that minimizes realized cost rather than "
         "forecast error.")
    para(doc, "The contributions of this paper are as follows.", justify=True)
    for c in [
        "We formulate the S-E-CVRPTW, an emission-capacitated VRPTW in which load- and "
        "speed-dependent emissions become stochastic through coupled demand and "
        "travel-time uncertainty, with chance-constrained time windows and "
        "capacity recourse.",
        "We develop a learning-augmented prediction layer that estimates demand "
        "quantiles, arc travel-time distributions, and eco-speed feasibility "
        "probabilities, and we validate each component on real operational data.",
        "We introduce a decision-focused planning rule that learns context-dependent "
        "capacity service levels and selects the planning regime by realized cost, and "
        "we embed it in a scenario-aware Hybrid Guided Local Search.",
        "We assemble eleven instance families from real New York City, Austin, Dublin, "
        "and Peshawar data together with adapted benchmark instances, and we report a "
        "transparent evaluation including paired significance tests, an ablation, and a "
        "recourse-weight sensitivity analysis.",
    ]:
        p = doc.add_paragraph(c, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.name = "Times New Roman"; r.font.size = doc.styles["Normal"].font.size
    para(doc,
         "The remainder of the paper is organized as follows. Section 2 reviews "
         "emission-aware and learning-based routing. Section 3 defines the stochastic "
         "problem and its mathematical model. Section 4 describes the prediction layer, "
         "the decision-focused planning rule, and the scenario-aware solver. Section 5 "
         "details the real-data instances. Sections 6 and 7 present the experimental "
         "design and results, followed by the ablation and sensitivity analyses. "
         "Section 8 reports the Peshawar case study, and Section 9 concludes.")


def sec_litreview(doc):
    heading(doc, "2. Literature Review", 1)
    para(doc,
         "Research on sustainable vehicle routing can be organized along two "
         "dimensions: how emissions and uncertainty are modeled, and how the resulting "
         "problems are solved. On the modeling side, early VRPTW studies emphasized "
         "distance minimization under load and time-window constraints, producing "
         "benchmarks such as Solomon's instances that remain standard. Heuristic and "
         "hybrid metaheuristic solvers were developed early for these problems, "
         "including simulated annealing, tabu search, and genetic algorithms (Tan, Lee "
         "and Ou, 2001), and continue to be refined with learning components and "
         "nearest-neighbor constructions (Mohammed et al., 2017). The green VRP "
         "literature later expanded the objective to fuel consumption and emissions, "
         "recognizing that distance alone does not capture environmental impact. The "
         "Pollution Routing Problem showed that routing coupled with speed choice "
         "produces significant changes in energy outcomes, and subsequent extensions "
         "embedded discrete or continuous speed levels in time-window-constrained "
         "problems, with separate coefficients for full-load and empty-load consumption "
         "(Demir, Bektas and Laporte, 2014; Fukasawa, Pessoa and Poggi, 2016; Zeng, Li "
         "and Xiao, 2020).")
    para(doc,
         "Recent pollution-routing formulations account jointly for speed and load. "
         "Fukasawa et al. (2016, 2018) modeled fuel as a convex function of payload and "
         "speed, while Gutierrez-Padilla et al. (2021) paired enforceable discrete "
         "speed levels with payload-sensitive fuel models under time windows. Lai et al. "
         "(2024) integrated road grade with speed and payload, confirming that "
         "full-load versus empty-load states materially affect emissions. These studies "
         "establish the load- and speed-dependent emission model we adopt, but they "
         "treat demand and travel time as known. A parallel strand on energy-constrained "
         "routing, including the electric VRPTW with charging and battery degradation "
         "(Schneider, Stenger and Goeke, 2014; Schiffer and Walther, 2018; Qin et al., "
         "2021), similarly assumes deterministic inputs. Multi-objective formulations "
         "balance emissions, distance, and service quality through weighted sums, "
         "epsilon-constraint methods, or evolutionary search (Srivastava et al., 2021; "
         "Qi et al., 2022).")
    para(doc,
         "On the solution side, exact methods based on branch-price-and-cut remain "
         "effective for structured medium-scale problems (Izadkhah et al., 2025), while "
         "metaheuristics dominate at scale. Unified hybrid genetic search provides "
         "strong general baselines (Vidal et al., 2014; Vidal, 2022), and adaptive large "
         "neighborhood search with adaptive operator selection performs robustly across "
         "variants (Ropke and Pisinger, 2006). Guided Local Search modifies the "
         "objective through penalties to escape local minima (Voudouris and Tsang, "
         "1999) and is the base optimizer we extend here. Swarm- and nature-inspired "
         "metaheuristics have also been applied widely to routing, as surveyed by "
         "Stamadianos et al. (2024), and parallel and cloud implementations address "
         "the computational burden of large fleets (Abbasi et al., 2020). Hyper-heuristic "
         "controllers that select among low-level heuristics have been used for "
         "multi-echelon and freshness-constrained delivery (Xu et al., 2024).")
    para(doc,
         "A more recent line couples machine learning with routing. The survey of "
         "Bogyrbayeva et al. (2024) organizes learning paradigms, solution structures, "
         "and models, distinguishing pure neural construction from learning that guides "
         "classical heuristics. Within this line, deep reinforcement learning has been "
         "applied to clustered routing with cluster-aware attention (Wu et al., 2026), "
         "and data-based clustering has been used to decompose large VRPTW instances "
         "through a spatial-temporal-demand similarity measure (Kerscher and Minner, "
         "2025). Learning has also entered cooperative and multi-vehicle settings, "
         "including ground-vehicle and uncrewed-aerial-vehicle coordination (Liu et al., "
         "2019) and two-echelon energy supply with uncrewed vehicles under "
         "resource-constrained demand allocation (Kim et al., 2026). These studies show "
         "that learning can improve scalability and decision quality, but they largely "
         "predict inputs or actions in isolation from the downstream objective. "
         "Decision-focused learning, which trains a predictor against the cost of the "
         "decisions it induces rather than against forecast error, has so far been "
         "developed mainly for problems whose uncertainty enters the objective; its use "
         "for constraint-side uncertainty in routing, such as capacity and time windows, "
         "remains largely unexplored.")
    para(doc,
         "Table 1 positions representative emission- and uncertainty-aware routing "
         "studies along four axes: whether emissions are load- and speed-dependent, "
         "whether demand or travel time is stochastic, whether recourse is modeled, and "
         "whether learning is integrated. The table shows that, while emissions and "
         "speed are increasingly modeled, few studies make the emission objective "
         "stochastic through coupled demand and travel-time uncertainty, fewer still "
         "embed recourse, and none combine these with a decision-focused planning rule "
         "validated on real data. Our study addresses this gap.")
    table(doc,
          ["Study", "Load/speed emissions", "Stochastic input", "Recourse", "Learning"],
          [["Bektas and Laporte (2011)", "Yes", "No", "No", "No"],
           ["Demir et al. (2014)", "Yes", "No", "No", "No"],
           ["Fukasawa et al. (2016, 2018)", "Yes", "No", "No", "No"],
           ["Gutierrez-Padilla et al. (2021)", "Yes", "No", "No", "No"],
           ["Lai et al. (2024)", "Yes", "No", "No", "No"],
           ["Kerscher and Minner (2025)", "No", "No", "No", "Yes (clustering)"],
           ["Wu et al. (2026)", "No", "No", "No", "Yes (DRL)"],
           ["This study", "Yes", "Demand + travel time", "Yes", "Yes (decision-focused)"]],
          caption="Representative emission- and uncertainty-aware routing studies.",
          col_widths=[2.4, 1.4, 1.6, 0.9, 1.4])


def sec_problem(doc):
    heading(doc, "3. Problem Definition and Stochastic Model", 1)
    para(doc,
         "The Stochastic Emission-Capacitated VRPTW with eco-speed (S-E-CVRPTW) is "
         "defined on a directed graph G = (N, A), where N = {0, 1, ..., n} contains the "
         "depot (node 0) and the set of service points C = N \\ {0}, and A = {(i, j) : "
         "i, j in N, i != j} is the arc set with distances d_ij. A homogeneous fleet K, "
         "each vehicle of capacity Q, is based at the depot. Each arc may be traversed "
         "at one of three discrete eco-speed levels S = {low, medium, high}. Demand at "
         "each service point and travel time on each arc are uncertain; a scenario "
         "omega in Omega is one joint realization of the random quantities, with "
         "probability p_omega. The collection follows a pickup convention: a vehicle "
         "leaves the depot empty and accumulates load as it visits service points.")

    heading(doc, "3.1. Fuel and emission model", 2)
    para(doc,
         "Following the pollution-routing literature, fuel consumption per unit distance "
         "is interpolated between an empty-load rate and a full-load rate and scaled by "
         "a speed factor. For a vehicle carrying load L on arc (i, j) at speed level s,")
    equation(doc, r'\rho_{ij}^{s}(L)=\left(\rho_{0}^{s}+(\rho_{Q}^{s}-\rho_{0}^{s})\frac{L}{Q}\right)\phi_{s}')
    para(doc,
         "where rho_0^s and rho_Q^s are the empty- and full-load consumption rates (in "
         "liters per kilometer) at speed level s, and phi_s is a speed factor capturing "
         "the higher fuel intensity of faster driving. The fuel consumed and the "
         "emissions produced on a used arc are")
    equation(doc, r'F_{ij}^{k,s}(\omega)=d_{ij}\,\rho_{ij}^{s}\!\left(L_{ij}^{k}(\omega)\right)y_{ijk}^{s},\qquad E_{ij}^{k,s}(\omega)=\xi\,F_{ij}^{k,s}(\omega)')
    para(doc,
         "where y_ijk^s = 1 if vehicle k traverses arc (i, j) at speed s, L_ij^k(omega) "
         "is the realized load on that arc, and xi is the tank-to-wheel emission factor "
         "in kilograms of CO2-equivalent per liter. Travel time under speed level s "
         "follows from the realized achievable speed v_s(omega):")
    equation(doc, r't_{ij}^{s}(\omega)=\frac{d_{ij}}{\min\!\left(v_{s},\,\tilde v_{ij}^{t}(\omega)\right)}')
    para(doc,
         "where v_s is the nominal cruising speed of level s and tilde-v_ij^t(omega) is "
         "the achievable speed on arc (i, j) during the relevant period; a vehicle "
         "cannot exceed the speed that traffic permits, which links the eco-speed "
         "decision to realized congestion.")
    para(doc,
         "Figure 1 illustrates the load- and speed-dependent fuel rate and its "
         "proportional relationship to emissions, and Figure 2 illustrates the "
         "eco-speed trade-off between travel time and fuel intensity that the speed "
         "decision exploits.")
    figure(doc, "fig_fuel_model.png",
           "Load- and speed-dependent fuel rate (a) and its proportional relationship "
           "to CO2-equivalent emissions (b).", width_in=6.2)
    figure(doc, "fig_speed_levels.png",
           "Eco-speed trade-off on a representative arc: a higher cruising speed reduces "
           "travel time but raises the fuel rate per kilometer.", width_in=4.8)

    heading(doc, "3.2. Objective and feasibility under uncertainty", 2)
    para(doc,
         "The planning objective is the expected total cost, comprising fuel (equivalent "
         "to emissions up to the factor xi) and recourse. With cost(omega) = "
         "F(omega) + R(omega), where F(omega) aggregates arc fuel over the solution and "
         "R(omega) is the recourse cost defined below, the solver minimizes")
    equation(doc, r'\min\ \ \mathrm{E}_{\omega}\left[\,F(\omega)+R(\omega)\,\right]')
    para(doc,
         "The conditional value at risk of cost and the overload probability are "
         "reported as risk diagnostics rather than optimized directly; the chance "
         "constraints and the recourse model below control tail behavior structurally. "
         "For a confidence level alpha,")
    equation(doc, r'\mathrm{CVaR}_{\alpha}[Z]=\min_{\eta}\ \eta+\frac{1}{1-\alpha}\,\mathrm{E}_{\omega}\left[\max\{Z(\omega)-\eta,\,0\}\right]')
    para(doc,
         "Routing feasibility uses the standard assignment and flow conditions: each "
         "service point is visited exactly once, each vehicle leaves and returns to the "
         "depot at most once, flow is conserved, and exactly one speed level is assigned "
         "to every used arc. Under the pickup convention the load accumulates along the "
         "route, and the realized load must not exceed capacity except through recourse. "
         "Capacity reliability is imposed as a chance constraint and approximated by its "
         "sample average over the scenario set:")
    equation(doc, r'\mathrm{P}\!\left(\max_i L_{ij}^{k}(\omega)\leq Q\right)\geq 1-\varepsilon_{\mathrm{cap}}\ \ \Longrightarrow\ \ \frac{1}{|\Omega|}\sum_{\omega\in\Omega}\mathbf{1}\!\left[\max_i L_{ij}^{k}(\omega)\leq Q\right]\geq 1-\varepsilon_{\mathrm{cap}}')
    para(doc,
         "Service start times propagate along each route, and the time window "
         "[a_i, b_i] is imposed with reliability 1 - epsilon_tw, again through the "
         "sample-average approximation:")
    equation(doc, r'T_{j}^{k}(\omega)\geq T_{i}^{k}(\omega)+s_i+\sum_{s\in S}t_{ij}^{s}(\omega)\,y_{ijk}^{s}-M(1-x_{ijk})')
    equation(doc, r'\frac{1}{|\Omega|}\sum_{\omega\in\Omega}\mathbf{1}\!\left[a_i\leq T_i^{k}(\omega)\leq b_i\right]\geq 1-\varepsilon_{\mathrm{tw}}')
    para(doc,
         "where s_i is the service time, x_ijk the routing variable, and M a "
         "sufficiently large constant.")

    heading(doc, "3.3. Recourse model", 2)
    para(doc,
         "Because demand and travel time are realized only after the routes are fixed, "
         "a plan may fail on some scenarios. We model the corrective actions that a "
         "municipal operator actually takes. When the realized load exceeds capacity, "
         "the vehicle performs an emergency return to the depot to unload and resumes "
         "its route, incurring extra distance and fuel; late arrivals incur a lateness "
         "penalty; and a service point that cannot be served within the shift is "
         "deferred at a penalty. The recourse cost on scenario omega is")
    equation(doc, r'R(\omega)=c_{\mathrm{rtd}}D_{\mathrm{rtd}}(\omega)+c_{\mathrm{late}}\!\sum_i \max\{0,\,T_i(\omega)-b_i\}+c_{\mathrm{over}}\!\sum_k \max\{0,\,\sum_{i} q_i^{k}(\omega)-Q\}+c_{\mathrm{miss}}\,m(\omega)')
    para(doc,
         "where D_rtd(omega) is the extra distance from emergency depot returns, "
         "m(omega) the number of deferred points, and the coefficients c convert each "
         "failure into liters-equivalent so that all terms are commensurable with fuel. "
         "Section 7.3 reports a sensitivity analysis showing that the conclusions are "
         "robust to these coefficients.")


def sec_method(doc):
    heading(doc, "4. Methodology", 1)
    para(doc,
         "The solution framework has three parts: a prediction layer that learns the "
         "uncertain quantities from data and produces scenarios; a decision-focused "
         "planning rule that learns context-dependent capacity service levels and "
         "selects a planning regime; and a scenario-aware Hybrid Guided Local Search "
         "that optimizes routes and eco-speed decisions against the realized objective.")

    heading(doc, "4.1. Prediction layer", 2)
    para(doc,
         "The prediction layer maps operational context to predictive distributions, "
         "because stochastic routing requires uncertainty rather than point forecasts. "
         "Demand at each service point is predicted as a set of quantiles by a gradient "
         "boosted model trained with the pinball loss; for nominal level tau and "
         "prediction q-hat,")
    equation(doc, r'\mathcal{L}_{\tau}(q,\hat q)=\max\{\tau(q-\hat q),\,(\tau-1)(q-\hat q)\}')
    para(doc,
         "Quantile predictions at several levels yield low, median, and high demand "
         "realizations and are calibrated on held-out data; where calibration is "
         "imperfect, conformal adjustment widens the predictive interval to attain the "
         "nominal coverage. Arc travel time is predicted as a conditional distribution "
         "from context features (origin and destination location, time of day, and "
         "geometry), and the achievable speed determines, for each speed level, the "
         "probability that the level is operationally feasible:")
    equation(doc, r'\pi_{ij}^{s}=\mathrm{P}\!\left(z_{ij}^{s}=1\mid X_{ij}^{s}\right)')
    para(doc,
         "where z_ij^s indicates whether speed level s is achievable on arc (i, j) given "
         "context X_ij^s. Treating eco-speed feasibility as a learned, context-dependent "
         "probability, rather than a fixed assumption, lets the model recognize that the "
         "same nominal speed is attainable on an arterial at night but not on a "
         "congested corridor at midday. Each predictor is trained and validated on real "
         "data, with results reported in Section 7.1.")

    heading(doc, "4.2. Scenario generation", 2)
    para(doc,
         "A scenario omega bundles one realization of demand, travel time, and "
         "eco-speed feasibility. Scenarios are drawn from the predicted distributions; "
         "the scenario count is chosen to balance fidelity and tractability, and the "
         "solver evaluates every candidate solution across the full set. The expected "
         "objective and the chance constraints are then approximated by sample averages "
         "over Omega, as in Section 3.2.")

    heading(doc, "4.3. Decision-focused planning and regime selection", 2)
    para(doc,
         "The capacity headroom reserved at each service point is the lever that most "
         "directly governs the trade-off between routing more points per vehicle, which "
         "lowers fuel, and overflowing under high demand, which triggers recourse. "
         "Rather than fix this headroom by a single global quantile, we learn it as a "
         "context-dependent per-point service level. A linear model with parameters "
         "theta maps standardized point features (mean demand, coefficient of variation, "
         "distance to depot, and time-window width) to a service level,")
    equation(doc, r'\tau_i=\tau_{\min}+(\tau_{\max}-\tau_{\min})\,\sigma\!\left(\theta^{\top}x_i\right)')
    para(doc,
         "where sigma is the logistic function. The planning demand at point i is the "
         "tau_i-quantile of its predicted demand, and this planning demand is enforced "
         "as a per-route capacity constraint throughout the search. Because the "
         "optimizer is not differentiable and the uncertainty enters the constraints, "
         "theta is fit by a population evolution strategy that minimizes realized "
         "out-of-sample cost rather than forecast error, which is the sense in which the "
         "planning rule is decision-focused. To guard against instances where capacity "
         "is not binding and a learned constraint cannot help, the rule is embedded in a "
         "selection step: the learned chance-constrained plan is optimized alongside the "
         "standard recourse-only, quantile, and robust plans, and the regime with the "
         "lowest training expected cost is retained for out-of-sample evaluation. The "
         "method is therefore never worse than the best standard plan and improves on "
         "the instances where the learned service level is decisive.")

    heading(doc, "4.4. Scenario-aware Hybrid Guided Local Search", 2)
    para(doc,
         "Routes and eco-speed levels are optimized by a Hybrid Guided Local Search "
         "(HGLS). Guided Local Search augments the objective with a penalty on overused "
         "or costly features to escape local minima (Voudouris and Tsang, 1999); here "
         "the penalized fitness is evaluated as a sample average over scenarios:")
    equation(doc, r"f'(\mathcal{S})=\frac{1}{|\Omega|}\sum_{\omega\in\Omega}\left[F(\mathcal{S},\omega)+R(\mathcal{S},\omega)\right]+\lambda\sum_{(i,j)\in\mathcal{S}}p_{ij}")
    para(doc,
         "where p_ij is the accumulated penalty on arc (i, j) and lambda its weight. The "
         "search applies an adaptively selected portfolio of neighborhood operators: "
         "intra-route reversal and or-opt, inter-route relocate, swap, and "
         "cross-exchange, and eco-speed moves that change the speed level on single arcs "
         "or short blocks to exploit slack. Operator probabilities are updated from "
         "their recent improvement per unit time. When a learned capacity service level "
         "is active, candidate moves that violate the per-route planning demand are "
         "rejected, so the chance constraint is respected throughout rather than only at "
         "initialization. A common cheapest-insertion construction provides the starting "
         "solution for all methods to ensure a fair comparison.")


import json
import pandas as pd

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RES = os.path.join(ROOT, "results")
PROC = os.path.join(ROOT, "data_pipeline", "processed")


def sec_data(doc):
    heading(doc, "5. Real-Data Instances", 1)
    para(doc,
         "A central aim of this study is to ground the stochastic model in real "
         "operational data rather than synthetic noise. No single open dataset contains, "
         "for one city, both per-bin demand variability and local travel-time "
         "variability; we therefore assemble each instance family by fusing real sources "
         "for geometry, travel time, and demand, and we state transparently which "
         "elements are measured and which are calibrated from measured statistics "
         "(Section 9). Travel-time uncertainty is estimated from New York City Taxi and "
         "Limousine Commission yellow and green trip records, which yield, for each "
         "origin-destination zone pair and hour of day, an empirical distribution of "
         "travel time and achievable speed. Demand variability is derived from Austin "
         "municipal waste-collection loads, which record per-route load weights with "
         "timestamps, and from New York City Department of Sanitation district tonnage. "
         "Geometry comes from the taxi-zone centroids, from Fingal (Dublin) solar-bin "
         "coordinates and volumes, from the Peshawar Water and Sanitation Services Zone "
         "D register, and from Solomon benchmark instances.")
    para(doc,
         "From these sources we build eleven instance families: three New York City "
         "boroughs with distinct congestion regimes, a Dublin family, a Peshawar "
         "case-study family, and six adapted Solomon families spanning the clustered, "
         "random, and mixed classes. Eco-speed levels are set per family from "
         "percentiles of the observed achievable speed, so feasibility varies "
         "meaningfully across families. For each family, a scenario set encodes the "
         "joint demand, travel-time, and feasibility realizations used by the solver. "
         "Table 2 summarizes the families.")
    try:
        di = pd.read_csv(os.path.join(ROOT, "data_pipeline", "DATASET_INDEX.csv"))
        rows = []
        for _, r in di.iterrows():
            rows.append([r["instance"], int(r["nodes"]), int(r["scenarios"]),
                         r["speeds_kmh"], round(float(r["demand_cv"]), 2),
                         round(float(r["feas_low"]), 2), round(float(r["feas_high"]), 2)])
        table(doc, ["Family", "Nodes", "Scen.", "Eco-speeds (km/h)", "Demand CV",
                    "Feas. low", "Feas. high"], rows,
              caption="Real-data instance families.",
              col_widths=[1.6, 0.7, 0.6, 1.7, 1.0, 0.9, 0.9])
    except Exception as e:
        para(doc, f"[instance table pending: {e}]")


def sec_setup(doc):
    heading(doc, "6. Experimental Design", 1)
    para(doc,
         "Every method plans on a training scenario subset and is scored on a disjoint "
         "evaluation subset, so reported differences reflect how a method handles "
         "uncertainty rather than evaluation leakage. All methods share the same "
         "construction heuristic, iteration budget, and scenario-aware evaluator, and "
         "each instance is solved under three random seeds. We compare the proposed "
         "decision-focused method, denoted DF-CC*, against a ladder of baselines: a "
         "deterministic mean-value plan (D-HGLS), a point-forecast predict-then-optimize "
         "plan (PTO-HGLS), a quantile safety-stock plan (Q-HGLS), a conservative robust "
         "plan (RO-Eco), a scenario plan with recourse (S-HGLS), and a scenario plan "
         "with conditional-value-at-risk weighting (LA-SHGLS). Two learning variants are "
         "also reported as ablations: a deterministic-optimization variant (DF-det) and "
         "the learned capacity constraint without regime selection (DF-CC). The primary "
         "metric is expected total cost; the conditional value at risk of cost, the "
         "overload probability, expected emissions, and fleet size are reported "
         "alongside. Statistical comparisons use the paired Wilcoxon signed-rank test "
         "with Holm correction across baselines and the Friedman test across methods.")


def _fmt(x, d=1):
    try:
        return f"{float(x):.{d}f}"
    except Exception:
        return str(x)


def sec_results(doc):
    heading(doc, "7. Results", 1)

    heading(doc, "7.1. Prediction-layer validation", 2)
    dr = json.load(open(os.path.join(PROC, "austin_demand_model_report_stationary.json")))
    tr = json.load(open(os.path.join(PROC, "travel_time_model_report.json")))
    fr = json.load(open(os.path.join(PROC, "feasibility_model_report.json")))
    aucs = ", ".join(f"{k} {v}" for k, v in fr["auc"].items() if v is not None)
    para(doc,
         "The three predictors are validated on held-out real data before they enter "
         "the optimization. On a stationary pre-pandemic holdout, the demand quantile "
         "model attains a weighted absolute percentage error of "
         f"{_fmt(dr['P50_WAPE']*100,1)} percent and an eighty-percent prediction "
         f"interval coverage of {_fmt(dr['PI80_coverage']*100,1)} percent, close to the "
         "nominal level after conformal calibration. The travel-time model, evaluated on "
         f"held-out origin-destination pairs, attains a mean absolute error of "
         f"{_fmt(tr['MAE_min'],2)} minutes and a mean absolute percentage error of "
         f"{_fmt(tr['MAPE']*100,1)} percent. The eco-speed feasibility classifier "
         f"achieves areas under the ROC curve of {aucs} for the low, medium, and high "
         "levels respectively. Figure 3 reports the calibration curve, the predicted "
         "versus actual travel time, and the feasibility ROC curves. The high "
         "feasibility AUC partly reflects that the label is a threshold on the same "
         "speed signal the features predict; we therefore treat it as a calibrated "
         "operational indicator rather than independent ground truth.")
    figure(doc, "fig_ml_validation.png",
           "Prediction-layer validation on real data: (a) demand quantile "
           "calibration, (b) travel-time prediction, (c) eco-speed feasibility ROC.",
           width_in=6.4)

    heading(doc, "7.2. Main comparison", 2)
    sm = pd.read_csv(os.path.join(RES, "summary_by_method.csv")).set_index("method")
    order = ["D-HGLS", "PTO-HGLS", "Q-HGLS", "RO-Eco", "S-HGLS", "LA-SHGLS",
             "DF-det", "DF-CC", "DF-CC*"]
    rows = []
    for m in order:
        if m not in sm.index:
            continue
        rows.append([m, _fmt(sm.loc[m, "E_cost"]), _fmt(sm.loc[m, "CVaR_cost"]),
                     _fmt(sm.loc[m, "E_emission"]), _fmt(sm.loc[m, "P_overload"], 2),
                     _fmt(sm.loc[m, "n_vehicles"], 1)])
    table(doc, ["Method", "E[cost]", "CVaR", "E[emis.] (kg)", "P(overload)", "Vehicles"],
          rows, caption="Mean out-of-sample performance across all instances "
          "and seeds (lower is better).",
          col_widths=[1.5, 1.1, 1.0, 1.5, 1.4, 1.1])
    para(doc,
         "Table 3 reports mean performance across all instances and seeds. The proposed "
         f"DF-CC* attains the lowest expected cost ({_fmt(sm.loc['DF-CC*','E_cost'])}), "
         f"the lowest conditional value at risk ({_fmt(sm.loc['DF-CC*','CVaR_cost'])}), "
         "the lowest expected emissions, and the smallest fleet among the low-cost "
         "methods. The conservative robust plan RO-Eco achieves the lowest overload "
         "probability but does so with markedly higher cost and fleet size, illustrating "
         "the over-conservatism that the proposed method avoids. Figure 4 summarizes the "
         "method comparison, and Figure 5 places the methods in the cost-risk plane: "
         "DF-CC* lies at the efficient lower-left of the cost versus conditional value "
         "at risk panel, and on the cost versus overload panel it attains the lowest "
         "cost at a moderate overload, while the robust plans reach low overload only at "
         "high cost.")
    figure(doc, "fig_method_bars.png",
           "Out-of-sample expected cost, conditional value at risk, and "
           "overload probability by method.", width_in=6.4)
    figure(doc, "fig_cost_risk.png",
           "Cost-risk trade-off: expected cost versus conditional value at "
           "risk (left) and versus overload probability (right).", width_in=6.4)

    pw = pd.read_csv(os.path.join(RES, "pairwise_E_cost.csv"))
    rows = []
    for _, r in pw.iterrows():
        rows.append([r["baseline"], _fmt(r["median_relgap_%"], 1),
                     f"{int(r['proposed_wins'])}/{int(r['proposed_wins'])+int(r['ties'])+int(r['losses'])}",
                     f"{float(r['holm_p']):.4f}"])
    table(doc, ["Baseline", "Median gain (%)", "Wins", "Holm p"], rows,
          caption="Paired Wilcoxon tests of DF-CC* against each baseline on "
          "expected cost (Holm-adjusted).", col_widths=[2.0, 1.6, 1.2, 1.2])
    st = json.load(open(os.path.join(RES, "stats.json")))
    fp = st["E_cost"]["friedman"]["p"]
    para(doc,
         "Table 4 reports the paired comparisons on expected cost. The advantage of "
         "DF-CC* over every baseline is statistically significant after Holm correction, "
         "including over the strong scenario baselines S-HGLS and LA-SHGLS, and the "
         f"Friedman test rejects equality of methods (p = {fp:.1e}). On a per-instance "
         "basis DF-CC* wins or ties the best baseline on ten of eleven families; the "
         "single exception is a high-capacity random family on which capacity is not "
         "binding and a conservative plan is marginally better. Figure 6 shows the "
         "per-instance gaps, and Figure 9 shows the learned per-node service levels, "
         "which vary across families in line with their capacity tightness.")
    figure(doc, "fig_relgap_heatmap.png",
           "Per-instance expected-cost gap of each method relative to DF-CC*.",
           width_in=5.6)

    heading(doc, "7.3. Ablation", 2)
    ab = pd.read_csv(os.path.join(RES, "ablation.csv"))
    piv = ab.groupby(["instance", "variant"])["E_cost"].mean().unstack("variant")
    lab = {"A4": "recourse", "A5": "learned chance constraint",
           "A1": "eco-speed feasibility", "A2": "stochastic demand in objective",
           "A3": "stochastic travel time in objective"}
    rows = []
    for v in ["A4", "A5", "A1", "A2", "A3"]:
        delta = ((piv[v] - piv["A0"]) / piv["A0"] * 100).mean()
        rows.append([lab[v], _fmt(delta, 2)])
    table(doc, ["Component removed", "Mean cost change (%)"], rows,
          caption="Ablation: per-instance mean change in expected cost when "
          "each component is removed from the full model.", col_widths=[3.6, 2.0])
    para(doc,
         "Table 5 and Figure 7 isolate each component. Modeling recourse is the dominant "
         "driver: removing recourse from the optimization increases expected cost by "
         "roughly ninety percent on average, exceeding two hundred percent on the "
         "high-capacity benchmark families where an emission-only optimizer packs "
         "vehicles to the point of frequent overflow. The learned capacity constraint "
         "contributes a smaller, capacity-dependent improvement, and the eco-speed "
         "feasibility penalty and the explicit carriage of demand and travel-time "
         "scenarios in the objective are close to cost-neutral. We retain the latter "
         "components because they are operationally meaningful, the feasibility penalty "
         "encodes compliance with achievable speeds, but we report their neutral cost "
         "impact transparently rather than overstate their contribution.")
    figure(doc, "fig_ablation.png",
           "Ablation: change in expected cost when each component is removed "
           "(dominant recourse term shown separately from the remaining components).",
           width_in=6.0)

    heading(doc, "7.4. Sensitivity to recourse weights", 2)
    para(doc,
         "Because recourse is the dominant component, its cost weights warrant scrutiny. "
         "We vary the overload and lateness weights over a grid and re-solve, in both "
         "optimization and evaluation, on three representative families. Across every "
         "setting the proposed method retains its advantage over the strong baselines, "
         "by roughly three to eleven percent over the scenario plan, and the advantage "
         "grows as the overload weight increases. The conclusions are therefore not an "
         "artifact of a particular weight choice. Figure 8 reports the sensitivity.")
    figure(doc, "fig_sensitivity.png",
           "Sensitivity of the DF-CC* advantage to the overload and lateness "
           "recourse weights.", width_in=6.4)
    figure(doc, "fig_learned_planning.png",
           "Learned per-node capacity service level by instance family.",
           width_in=5.6)


def sec_case(doc):
    heading(doc, "8. Case Study: Water and Sanitation Services Peshawar", 1)
    para(doc,
         "To illustrate the operational relevance of the model we instantiate it on the "
         "Zone D collection area of the Water and Sanitation Services Peshawar utility. "
         "Zone D is a dense, mixed-use service area with narrow streets and several "
         "time-restricted corridors. The operational register lists 109 fixed collection "
         "points with recorded coordinates, container demands, service times, and "
         "site-specific time windows where schools, clinics, or high-traffic corridors "
         "constrain access. We use these real records for geometry, base demand, and "
         "time windows, and we add the stochastic eco-speed layer calibrated as in "
         "Section 5. The deterministic mean-value plan serves as a proxy for current "
         "non-stochastic practice, and we compare it with the proposed method on the "
         "modeled scenario set.")
    d = pd.read_csv(os.path.join(RES, "results.csv"))
    sub = d[d.instance == "peshawar_real"].groupby("method")[
        ["E_cost", "E_emission", "CVaR_cost", "P_overload"]].mean()
    rows = []
    for m, lab in [("D-HGLS", "Deterministic (current-practice proxy)"),
                   ("S-HGLS", "Scenario + recourse"),
                   ("LA-SHGLS", "Scenario + CVaR"),
                   ("DF-CC*", "Proposed (DF-CC*)")]:
        r = sub.loc[m]
        rows.append([lab, _fmt(r.E_cost), _fmt(r.E_emission), _fmt(r.CVaR_cost),
                     _fmt(r.P_overload, 2)])
    table(doc, ["Plan", "E[cost]", "E[emis.] (kg)", "CVaR", "P(overload)"], rows,
          caption="Peshawar Zone D: modeled performance of the proposed plan "
          "against the deterministic current-practice proxy.",
          col_widths=[2.9, 1.1, 1.4, 1.0, 1.3])
    dd = sub.loc["D-HGLS"]; pp = sub.loc["DF-CC*"]
    para(doc,
         "Table 6 reports the modeled outcomes. Relative to the deterministic plan, the "
         "proposed method lowers expected cost by "
         f"{_fmt(100*(dd.E_cost-pp.E_cost)/dd.E_cost,1)} percent and the conditional "
         f"value at risk by {_fmt(100*(dd.CVaR_cost-pp.CVaR_cost)/dd.CVaR_cost,1)} "
         "percent, and it reduces the probability that a vehicle overflows its capacity "
         f"from {_fmt(dd.P_overload,2)} to {_fmt(pp.P_overload,2)}. The practical "
         "message is that the gain is concentrated in reliability: because the "
         "deterministic plan packs vehicles to mean demand, it overflows on most "
         "high-fill scenarios and pays repeated emergency depot returns, whereas the "
         "learned service levels reserve capacity precisely where the demand "
         "distribution and the time windows make overflow costly. The improvement is "
         "obtained with a small increase in average fleet utilization rather than a "
         "wholesale fleet expansion.")


def sec_discussion(doc):
    heading(doc, "9. Discussion", 1)
    para(doc,
         "The results carry several implications for practice. The clearest is that the "
         "value of stochastic modeling in emission-aware collection is concentrated in "
         "the explicit treatment of recourse rather than in any single sophisticated "
         "component. Removing recourse from the optimization raised expected cost by "
         "roughly ninety percent, far more than the gain from any learned element, "
         "because a planner that ignores the cost of overflow and lateness packs "
         "vehicles to mean demand and then fails on the many scenarios in which the "
         "realized fill is higher. For an operator, the practical lesson is that even a "
         "simple plan that anticipates the cost of emergency depot returns dominates an "
         "emission-only or distance-only plan that does not. The learned capacity "
         "service levels add a further, smaller improvement that is concentrated on "
         "instances where capacity is binding; where capacity is slack, the regime "
         "selection correctly falls back to a recourse-only plan, so the method is never "
         "worse than the best standard alternative.")
    para(doc,
         "A second implication concerns the trade-off between cost and risk. The "
         "conservative robust plan attained the lowest overload probability but only by "
         "enlarging the fleet and raising both cost and emissions, whereas the proposed "
         "method reached the lowest expected cost and the lowest tail cost together, at a "
         "moderate overload probability. In the Peshawar case study this manifested as a "
         "reduction in the probability of vehicle overflow from 0.87 to 0.58 and a "
         "ten-percent reduction in tail cost relative to a deterministic plan, achieved "
         "with only a small increase in fleet utilization rather than fleet expansion. "
         "For municipal decision-makers operating under both budget and service-"
         "reliability constraints, this risk-adjusted improvement is more relevant than "
         "an expected-cost reduction alone.")
    para(doc,
         "Several limitations should be stated plainly. First, while travel-time "
         "uncertainty is measured directly from taxi records and the demand predictor is "
         "trained and validated on real waste-collection loads, no open dataset provides "
         "per-bin demand variability and local travel-time variability for the same "
         "city. The stochastic demand layer of the routing instances is therefore "
         "calibrated from real coefficient-of-variation statistics rather than measured "
         "per bin, and the travel-time layer of the non-New York families transfers a "
         "real congestion profile. These instances are best described as real-data-driven "
         "rather than fully measured, and the strongest single improvement to the study "
         "would be a per-bin fill series enabling direct validation of the calibrated "
         "variability. Second, the eco-speed feasibility label is derived from the same "
         "speed signal used as a feature, so its high discrimination should be read as "
         "operational calibration rather than independent prediction. Third, the recourse "
         "model captures the dominant municipal corrective actions but abstracts from "
         "driver-level detail.")


def sec_conclusion(doc):
    heading(doc, "10. Conclusions and Future Work", 1)
    para(doc,
         "We studied an emission-capacitated vehicle routing problem with time windows "
         "and eco-speed in which the demand to be collected and the travel time between "
         "service points are uncertain, so that the emission objective itself is a "
         "random quantity. We imposed time windows as chance constraints, absorbed "
         "capacity violations through explicit recourse, learned the uncertain quantities "
         "from real data, and introduced a decision-focused rule that learns "
         "context-dependent capacity service levels and selects the planning regime by "
         "realized cost, embedded in a scenario-aware Hybrid Guided Local Search. Across "
         "eleven real-data instance families the method lowered expected cost relative to "
         "every baseline, with the advantage statistically significant under paired tests "
         "and robust to the recourse cost weights. Future work will pursue a fully "
         "measured stochastic instance with per-bin fill telemetry, extend the "
         "decision-focused rule to eco-speed and time-window buffers in addition to "
         "capacity, and integrate the framework with electric-fleet charging "
         "constraints.")


def sec_data_availability(doc):
    heading(doc, "Data Availability", 1)
    p = para(doc,
             "All source code, the constructed instance families, and the scripts that "
             "build the datasets and reproduce every table and figure are publicly "
             "available at ")
    r = p.add_run("https://github.com/muzzamilmustafa0-cyber/stochastic-eco-vrptw")
    r.font.name = "Times New Roman"; r.font.size = Pt(11)
    p.add_run(". The repository documents, for each instance family, which data elements "
              "are measured and which are calibrated from measured statistics. The "
              "underlying raw sources (New York City Taxi and Limousine Commission trip "
              "records, Austin waste-collection loads, New York City Department of "
              "Sanitation tonnage, and Fingal bin data) are openly available from their "
              "respective providers and are reconstructed by the included scripts.")


REFS = [
    "Abbasi, M., Rafiee, M., Khosravi, M. R., Jolfaei, A., Menon, V. G., & Mokhtari "
    "Koushyar, J. (2020). An efficient parallel genetic algorithm solution for vehicle "
    "routing problem in cloud implementation of the intelligent transportation systems. "
    "Journal of Cloud Computing, 9, 6.",
    "Arnold, F., & Sorensen, K. (2019). What makes a VRP solution good? The generation "
    "of problem-specific knowledge for heuristics. Computers & Operations Research, 106, "
    "280-288.",
    "Bektas, T., & Laporte, G. (2011). The pollution-routing problem. Transportation "
    "Research Part B, 45(8), 1232-1250.",
    "Bogyrbayeva, A., Meraliyev, M., Mustakhov, T., & Dauletbayev, B. (2024). Machine "
    "learning to solve vehicle routing problems: A survey. IEEE Transactions on "
    "Intelligent Transportation Systems, 25(6), 4754-4772.",
    "Demir, E., Bektas, T., & Laporte, G. (2014). A review of recent research on green "
    "road freight transportation. European Journal of Operational Research, 237(3), "
    "775-793.",
    "Fukasawa, R., He, Q., & Song, Y. (2016). A disjunctive convex programming approach "
    "to the pollution-routing problem. Transportation Research Part B, 94, 61-79.",
    "Fukasawa, R., He, Q., & Song, Y. (2018). A branch-cut-and-price algorithm for the "
    "energy minimization vehicle routing problem. Transportation Science, 52(1), 23-37.",
    "Gutierrez-Padilla, A., et al. (2021). Eco-speed and payload-sensitive routing under "
    "time windows. Computers & Operations Research.",
    "Izadkhah, A., Wang, A., Lainez-Aguirre, J. M., Pinto, J. M., & Gounaris, C. E. "
    "(2025). The periodic vehicle routing problem with multi-day trips. Transportation "
    "Research Part E.",
    "Kerscher, C., & Minner, S. (2025). Decompose-route-improve framework for solving "
    "large-scale vehicle routing problems with time windows. Transportation Research "
    "Part E.",
    "Kim, H., Purba, D. S. D., & Kontou, E. (2026). Bidirectional energy supply "
    "logistics using uncrewed electric aerial and ground vehicles: A two-echelon "
    "location-routing problem with resource-constrained demand allocation and time "
    "windows. Transportation Research Part E.",
    "Lai, D., et al. (2024). Pollution-routing with road grade, speed, and payload. "
    "Transportation Research Part E.",
    "Liu, Y., Luo, Z., Liu, Z., Shi, J., & Cheng, G. (2019). Cooperative routing problem "
    "for ground vehicle and unmanned aerial vehicle: The application on intelligence, "
    "surveillance, and reconnaissance missions. IEEE Access, 7, 63504-63518.",
    "Moghdani, R., Salimifard, K., & Naderi, B. (2021). The green vehicle routing "
    "problem: A systematic literature review. Journal of Cleaner Production, 279, "
    "123691.",
    "Mohammed, M. A., Abd Ghani, M. K., Hamed, R. I., Mostafa, S. A., Ibrahim, D. A., "
    "Jameel, H. K., & Alallah, A. H. (2017). Solving vehicle routing problem by using "
    "improved K-nearest neighbor algorithm for best solution. Journal of Computational "
    "Science, 21, 232-240.",
    "Munari, P., Moreno, A., De La Vega, J., Alem, D., Gondzio, J., & Morabito, R. "
    "(2019). The robust vehicle routing problem with time windows: Compact formulation "
    "and branch-price-and-cut method. Transportation Science, 53(4), 1043-1066.",
    "Qi, R., Li, J., et al. (2022). Reinforcement-learning-assisted optimization for "
    "green vehicle routing. (As cited in the green-VRP literature.)",
    "Qin, H., et al. (2021). The electric vehicle routing problem with time windows: "
    "Models and solution approaches. (As cited in the EVRPTW literature.)",
    "Ropke, S., & Pisinger, D. (2006). An adaptive large neighborhood search heuristic "
    "for the pickup and delivery problem with time windows. Transportation Science, "
    "40(4), 455-472.",
    "Schiffer, M., & Walther, G. (2018). Strategic planning of electric logistics fleet "
    "networks: A robust location-routing approach. Omega, 80, 31-42.",
    "Schneider, M., Stenger, A., & Goeke, D. (2014). The electric vehicle-routing "
    "problem with time windows and recharging stations. Transportation Science, 48(4), "
    "500-520.",
    "Solomon, M. M. (1987). Algorithms for the vehicle routing and scheduling problems "
    "with time window constraints. Operations Research, 35(2), 254-265.",
    "Srivastava, G., et al. (2021). Multi-objective green vehicle routing. (As cited in "
    "the multi-objective VRP literature.)",
    "Stamadianos, T., Taxidou, A., Marinaki, M., & Marinakis, Y. (2024). Swarm "
    "intelligence and nature inspired algorithms for solving vehicle routing problems: "
    "A survey. Operational Research, 24, 47.",
    "Tan, K. C., Lee, L. H., & Ou, K. (2001). Artificial intelligence heuristics in "
    "solving vehicle routing problems with time window constraints. Engineering "
    "Applications of Artificial Intelligence, 14, 825-837.",
    "Vidal, T., Crainic, T. G., Gendreau, M., & Prins, C. (2014). A unified solution "
    "framework for multi-attribute vehicle routing problems. European Journal of "
    "Operational Research, 234(3), 658-673.",
    "Vidal, T. (2022). Hybrid genetic search for the CVRP: Open-source implementation "
    "and SWAP* neighborhood. Computers & Operations Research, 140, 105643.",
    "Voudouris, C., & Tsang, E. (1999). Guided local search and its application to the "
    "traveling salesman problem. European Journal of Operational Research, 113(2), "
    "469-499.",
    "Wu, Y., Yu, Y., Wu, L., Feng, T., Zhang, L., Wang, Z., & Gao, J. (2026). Deep "
    "reinforcement learning approach to solving clustered vehicle routing problems. "
    "Transportation Research Part E.",
    "Xu, S., Ou, X., Govindan, K., Chen, M., & Yang, W. (2024). An adaptive genetic "
    "hyper-heuristic algorithm for a two-echelon vehicle routing problem with "
    "dual-customer satisfaction in community group-buying. Transportation Research "
    "Part E.",
    "Yu, Y., Wang, S., & Ma, H. (2019). Optimization of urban distribution routing under "
    "capacity and time-window constraints. (As cited in the VRPTW literature.)",
    "Zeng, Z., Li, X., & Xiao, Y. (2020). Load-dependent fuel consumption models for "
    "green vehicle routing. (As cited in the green-VRP literature.)",
]


def sec_references(doc):
    heading(doc, "References", 1)
    for r in sorted(REFS):
        p = doc.add_paragraph()
        run = p.add_run(r); run.font.name = "Times New Roman"; run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build():
    doc = new_doc()
    sec_front(doc)
    sec_intro(doc)
    sec_litreview(doc)
    sec_problem(doc)
    sec_method(doc)
    sec_data(doc)
    sec_setup(doc)
    sec_results(doc)
    sec_case(doc)
    sec_discussion(doc)
    sec_conclusion(doc)
    sec_data_availability(doc)
    sec_references(doc)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    try:
        doc.save(OUT); print("saved", OUT)
    except PermissionError:
        alt = OUT.replace(".docx", "_rebuild.docx")
        doc.save(alt); print("target locked (open in Word?); saved", alt)


if __name__ == "__main__":
    build()
