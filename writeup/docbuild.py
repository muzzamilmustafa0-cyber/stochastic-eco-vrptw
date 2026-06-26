"""Shared builder helpers for the manuscript: page setup, styles, headings,
paragraphs, numbered equations (rendered with matplotlib), figures, and tables."""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
EQDIR = os.path.join(HERE, "eq"); os.makedirs(EQDIR, exist_ok=True)
FIGDIR = os.path.join(os.path.dirname(HERE), "results", "figures")

_eqn = [0]


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.27), Inches(11.69)   # A4
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(1.0)
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"; st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
    return doc


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"
    r.font.size = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    return p


def para(doc, text, justify=True, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.bold = bold
    r.font.name = "Times New Roman"; r.font.size = Pt(size)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def render_eq(latex, fname, fontsize=15):
    path = os.path.join(EQDIR, fname)
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.text(0, 0, f"${latex}$", fontsize=fontsize)
    fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.04, transparent=True)
    plt.close(fig)
    return path


def equation(doc, latex, width_in=None):
    """Centered numbered display equation."""
    _eqn[0] += 1; n = _eqn[0]
    img = render_eq(latex, f"eq{n}.png")
    from PIL import Image
    w, h = Image.open(img).size
    win = (w / 220) if width_in is None else width_in
    win = min(win, 5.6)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Inches(5.6); tbl.columns[1].width = Inches(0.7)
    c0, c1 = tbl.rows[0].cells
    pc = c0.paragraphs[0]; pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.add_run().add_picture(img, width=Inches(win))
    pn = c1.paragraphs[0]; pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rn = pn.add_run(f"({n})"); rn.font.name = "Times New Roman"; rn.font.size = Pt(11)
    _no_borders(tbl)
    return n


def figure(doc, fname, caption, width_in=6.0):
    path = os.path.join(FIGDIR, fname)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"
    cap.paragraph_format.space_after = Pt(10)


def table(doc, header, rows, caption=None, col_widths=None, fontsize=9):
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(10); r.font.name = "Times New Roman"
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; c.paragraphs[0].clear()
        rr = c.paragraphs[0].add_run(str(h)); rr.bold = True; rr.font.size = Pt(fontsize)
        rr.font.name = "Times New Roman"
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].paragraphs[0].clear()
            rr = cells[j].paragraphs[0].add_run(str(v)); rr.font.size = Pt(fontsize)
            rr.font.name = "Times New Roman"
    if col_widths:
        for j, w in enumerate(col_widths):
            for r_ in t.rows:
                r_.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); borders.append(e)
    tblPr.append(borders)
